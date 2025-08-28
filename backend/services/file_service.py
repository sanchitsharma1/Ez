import asyncio
import logging
import os
import hashlib
from typing import Optional, Dict, Any, List
from pathlib import Path
import aiofiles
from datetime import datetime

# Document processing imports
import pypdf
from docx import Document
import unstructured
from unstructured.partition.auto import partition

from core.config import settings

logger = logging.getLogger(__name__)

class FileService:
    """Service for handling file uploads, processing, and text extraction"""
    
    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.max_file_size = settings.MAX_FILE_SIZE
        self.allowed_extensions = settings.ALLOWED_FILE_TYPES
        
        # Ensure upload directory exists
        self.upload_dir.mkdir(exist_ok=True)
    
    async def initialize(self):
        """Initialize file service"""
        try:
            # Create subdirectories
            (self.upload_dir / "documents").mkdir(exist_ok=True)
            (self.upload_dir / "audio").mkdir(exist_ok=True)
            (self.upload_dir / "temp").mkdir(exist_ok=True)
            
            logger.info("File service initialized successfully")
            
        except Exception as e:
            logger.error(f"Failed to initialize file service: {e}")
            raise
    
    async def save_uploaded_file(self, file_content: bytes, filename: str, content_type: str) -> Dict[str, Any]:
        """Save uploaded file and return file info"""
        try:
            # Validate file
            validation_result = self._validate_file(file_content, filename, content_type)
            if not validation_result["valid"]:
                return {"success": False, "error": validation_result["error"]}
            
            # Generate unique filename
            file_hash = hashlib.md5(file_content).hexdigest()
            file_ext = Path(filename).suffix.lower()
            unique_filename = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{file_hash[:8]}{file_ext}"
            
            # Determine subdirectory based on file type
            if file_ext in [".pdf", ".txt", ".docx", ".doc", ".rtf"]:
                subdir = "documents"
            elif file_ext in [".wav", ".mp3", ".mp4", ".flac", ".ogg"]:
                subdir = "audio"
            else:
                subdir = "temp"
            
            file_path = self.upload_dir / subdir / unique_filename
            
            # Save file
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(file_content)
            
            # Create file metadata
            file_info = {
                "success": True,
                "filename": filename,
                "unique_filename": unique_filename,
                "file_path": str(file_path),
                "file_size": len(file_content),
                "content_type": content_type,
                "file_hash": file_hash,
                "uploaded_at": datetime.utcnow().isoformat(),
                "category": subdir
            }
            
            logger.info(f"File saved successfully: {filename} -> {unique_filename}")
            return file_info
            
        except Exception as e:
            logger.error(f"Error saving uploaded file {filename}: {e}")
            return {"success": False, "error": str(e)}
    
    def _validate_file(self, file_content: bytes, filename: str, content_type: str) -> Dict[str, Any]:
        """Validate uploaded file"""
        try:
            # Check file size
            if len(file_content) > self.max_file_size:
                return {
                    "valid": False,
                    "error": f"File size ({len(file_content)} bytes) exceeds maximum allowed ({self.max_file_size} bytes)"
                }
            
            # Check file extension
            file_ext = Path(filename).suffix.lower()
            if file_ext not in self.allowed_extensions:
                return {
                    "valid": False,
                    "error": f"File type {file_ext} not allowed. Allowed types: {', '.join(self.allowed_extensions)}"
                }
            
            # Basic content validation
            if len(file_content) == 0:
                return {"valid": False, "error": "File is empty"}
            
            # Check for potentially malicious content
            if self._scan_for_malware(file_content):
                return {"valid": False, "error": "File failed security scan"}
            
            return {"valid": True, "error": None}
            
        except Exception as e:
            logger.error(f"Error validating file: {e}")
            return {"valid": False, "error": f"Validation failed: {str(e)}"}
    
    def _scan_for_malware(self, file_content: bytes) -> bool:
        """Basic malware scanning (placeholder for real scanner)"""
        try:
            # Check for suspicious patterns
            suspicious_patterns = [
                b"<script",
                b"javascript:",
                b"vbscript:",
                b"onload=",
                b"onerror="
            ]
            
            content_lower = file_content.lower()
            return any(pattern in content_lower for pattern in suspicious_patterns)
            
        except Exception:
            return True  # Err on the side of caution
    
    async def extract_text(self, file_path: str, file_type: str) -> Optional[str]:
        """Extract text from various file formats"""
        try:
            file_path_obj = Path(file_path)
            
            if not file_path_obj.exists():
                logger.error(f"File not found: {file_path}")
                return None
            
            file_ext = file_path_obj.suffix.lower()
            
            # Route to appropriate extraction method
            if file_ext == ".pdf":
                return await self._extract_pdf_text(file_path)
            elif file_ext in [".docx", ".doc"]:
                return await self._extract_docx_text(file_path)
            elif file_ext in [".txt", ".rtf"]:
                return await self._extract_plain_text(file_path)
            else:
                # Try unstructured as fallback
                return await self._extract_with_unstructured(file_path)
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return None
    
    async def _extract_pdf_text(self, file_path: str) -> Optional[str]:
        """Extract text from PDF files"""
        try:
            def extract_sync():
                with open(file_path, 'rb') as file:
                    reader = pypdf.PdfReader(file)
                    text_parts = []
                    
                    for page_num, page in enumerate(reader.pages):
                        try:
                            text = page.extract_text()
                            if text.strip():
                                text_parts.append(f"--- Page {page_num + 1} ---\n{text}\n")
                        except Exception as e:
                            logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                            continue
                    
                    return "\n".join(text_parts) if text_parts else None
            
            # Run in executor to avoid blocking
            return await asyncio.get_event_loop().run_in_executor(None, extract_sync)
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return None
    
    async def _extract_docx_text(self, file_path: str) -> Optional[str]:
        """Extract text from Word documents"""
        try:
            def extract_sync():
                doc = Document(file_path)
                text_parts = []
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text for cell in row.cells])
                        if row_text.strip():
                            text_parts.append(row_text)
                
                return "\n\n".join(text_parts) if text_parts else None
            
            return await asyncio.get_event_loop().run_in_executor(None, extract_sync)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return None
    
    async def _extract_plain_text(self, file_path: str) -> Optional[str]:
        """Extract text from plain text files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
                return content if content.strip() else None
                
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin1', 'cp1252', 'iso-8859-1']:
                try:
                    async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
                        content = await f.read()
                        return content if content.strip() else None
                except Exception:
                    continue
            
            logger.error(f"Could not decode text file: {file_path}")
            return None
            
        except Exception as e:
            logger.error(f"Error extracting plain text: {e}")
            return None
    
    async def _extract_with_unstructured(self, file_path: str) -> Optional[str]:
        """Extract text using unstructured library as fallback"""
        try:
            def extract_sync():
                elements = partition(filename=file_path)
                text_parts = []
                
                for element in elements:
                    if hasattr(element, 'text') and element.text.strip():
                        text_parts.append(element.text)
                
                return "\n\n".join(text_parts) if text_parts else None
            
            return await asyncio.get_event_loop().run_in_executor(None, extract_sync)
            
        except Exception as e:
            logger.error(f"Error with unstructured extraction: {e}")
            return None
    
    async def get_file_info(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Get information about a file"""
        try:
            file_path_obj = Path(file_path)
            
            if not file_path_obj.exists():
                return None
            
            stat_info = file_path_obj.stat()
            
            return {
                "filename": file_path_obj.name,
                "file_path": str(file_path_obj),
                "file_size": stat_info.st_size,
                "created_at": datetime.fromtimestamp(stat_info.st_ctime).isoformat(),
                "modified_at": datetime.fromtimestamp(stat_info.st_mtime).isoformat(),
                "file_extension": file_path_obj.suffix.lower(),
                "is_readable": os.access(file_path, os.R_OK)
            }
            
        except Exception as e:
            logger.error(f"Error getting file info for {file_path}: {e}")
            return None
    
    async def delete_file(self, file_path: str) -> bool:
        """Delete a file"""
        try:
            file_path_obj = Path(file_path)
            
            if file_path_obj.exists() and file_path_obj.is_file():
                file_path_obj.unlink()
                logger.info(f"File deleted: {file_path}")
                return True
            else:
                logger.warning(f"File not found for deletion: {file_path}")
                return False
                
        except Exception as e:
            logger.error(f"Error deleting file {file_path}: {e}")
            return False
    
    async def cleanup_temp_files(self, max_age_hours: int = 24):
        """Clean up temporary files older than specified hours"""
        try:
            temp_dir = self.upload_dir / "temp"
            if not temp_dir.exists():
                return
            
            current_time = datetime.now()
            deleted_count = 0
            
            for file_path in temp_dir.iterdir():
                if file_path.is_file():
                    file_age = current_time - datetime.fromtimestamp(file_path.stat().st_mtime)
                    
                    if file_age.total_seconds() > max_age_hours * 3600:
                        try:
                            file_path.unlink()
                            deleted_count += 1
                        except Exception as e:
                            logger.warning(f"Could not delete temp file {file_path}: {e}")
            
            if deleted_count > 0:
                logger.info(f"Cleaned up {deleted_count} temporary files")
                
        except Exception as e:
            logger.error(f"Error during temp file cleanup: {e}")
    
    async def get_storage_stats(self) -> Dict[str, Any]:
        """Get storage usage statistics"""
        try:
            stats = {
                "total_files": 0,
                "total_size": 0,
                "by_category": {}
            }
            
            for category in ["documents", "audio", "temp"]:
                category_dir = self.upload_dir / category
                if category_dir.exists():
                    file_count = 0
                    total_size = 0
                    
                    for file_path in category_dir.iterdir():
                        if file_path.is_file():
                            file_count += 1
                            total_size += file_path.stat().st_size
                    
                    stats["by_category"][category] = {
                        "file_count": file_count,
                        "total_size": total_size
                    }
                    
                    stats["total_files"] += file_count
                    stats["total_size"] += total_size
            
            return stats
            
        except Exception as e:
            logger.error(f"Error getting storage stats: {e}")
            return {
                "total_files": 0,
                "total_size": 0,
                "by_category": {}
            }